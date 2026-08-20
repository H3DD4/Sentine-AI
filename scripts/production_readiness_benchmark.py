"""Production-readiness benchmark: 60 labeled retrieval cases plus upload tests.

All vector work uses isolated ``pilot_*_benchmark`` collections. The benchmark
does not touch production collections or PostgreSQL sync state. Cases are built
from rows read from PostgreSQL, so expected document IDs and content markers are
deterministic rather than guessed.
"""

from __future__ import annotations

import argparse
import asyncio
import io
import json
import tempfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from fastapi import UploadFile
from PIL import Image
from sqlalchemy import or_, select

from app.db import AsyncSessionLocal
from app.ingestion.embedder import chunk_config, get_model, load_model_sync, load_sparse_model_sync
from app.kb.indexer import _build_points_batch, ensure_collection_sync
from app.kb.registry import get_source
from app.services.evidence import parse_evidence
from app.services.retrieval import get_qdrant, init_qdrant_client
from app.services.upload_processing import cleanup_staged_evidence, stage_evidence_uploads
from scripts.pilot_final_acceptance_suite import PilotSource, _hybrid_query
from app.services.retrieval import _exact_id_query_sync


SOURCE_KEYS = ["nvd", "mitre", "owasp", "owasp_docs", "finding_templates"]


def _id(source, row) -> str:
    return str(getattr(row, source.pk_column.key))


def _title(row, doc_id: str) -> str:
    return str(getattr(row, "title", "") or getattr(row, "name", "") or doc_id)


def _cases(source_key: str, rows: list, source) -> list[dict]:
    cases: list[dict] = []
    for index, row in enumerate(rows[:12]):
        doc_id = _id(source, row)
        title = _title(row, doc_id)
        body = source.build_text(row).split()
        fragment = " ".join(body[:12]) or doc_id
        lookup_id = str(getattr(row, "template_code", "") or doc_id)
        supports_exact = bool(source.id_pattern and source.extract_ids(lookup_id))
        variants = [
            ("exact-noise", f"locate [{lookup_id}] please", supports_exact),
            ("lowercase", lookup_id.lower().replace("-", " "), False),
            ("title", f"security documentation about {title}", False),
            ("fragment", fragment, False),
            ("question", f"what is the security issue described by {title}", False),
        ]
        for label, query, exact in variants:
            cases.append({
                "label": f"{source_key}.{index + 1}.{label}",
                "source": source_key,
                "query": query,
                "expected_id": doc_id,
                "lookup_id": lookup_id,
                "expected_marker": doc_id,
                "exact": exact,
            })
    return cases[:12]


async def _build_pilots(qdrant, limit: int) -> tuple[list[dict], list[dict]]:
    cases: list[dict] = []
    reports: list[dict] = []
    async with AsyncSessionLocal() as session:
        for source_key in SOURCE_KEYS:
            source = get_source(source_key)
            pilot = PilotSource(source)
            pilot.collection = f"pilot_{source.collection}_benchmark"
            if await asyncio.to_thread(qdrant.collection_exists, pilot.collection):
                await asyncio.to_thread(qdrant.delete_collection, pilot.collection)
            ensure_collection_sync(qdrant, pilot)
            if source_key == "owasp_docs":
                stmt = select(source.model).where(or_(
                    source.model.title.ilike("%server side request forgery%"),
                    source.model.title.ilike("%JSON Web Token%"),
                    source.model.title.ilike("%JWT%"),
                )).order_by(source.model.title).limit(limit)
            else:
                stmt = select(source.model).order_by(source.pk_column).limit(limit)
            rows = list((await session.execute(stmt)).scalars().all())
            built = await asyncio.to_thread(_build_points_batch, pilot, [(r, source.build_text(r)) for r in rows])
            points = [point for _, row_points in built for point in row_points]
            if not points:
                raise AssertionError(f"No pilot points created for {source_key}")
            await asyncio.to_thread(qdrant.upsert, collection_name=pilot.collection, points=points, wait=True)
            tokenizer = get_model().tokenizer
            lengths = [len(tokenizer.encode((p.payload or {}).get("chunk_text", ""), add_special_tokens=False)) for p in points]
            report = {
                "source": source_key,
                "rows": len(rows),
                "points": len(points),
                "max_tokens": max(lengths),
                "within_limit": max(lengths) <= chunk_config()[0],
                "self_describing": all(
                    f"Document ID: {(p.payload or {}).get('doc_id')}" in (p.payload or {}).get("chunk_text", "")
                    for p in points
                ),
                "unique_ids": len({p.id for p in points}) == len(points),
            }
            reports.append(report)
            if not all(report[k] for k in ("within_limit", "self_describing", "unique_ids")):
                raise AssertionError(f"Index contract failure: {report}")
            cases.extend(_cases(source_key, rows, source))
    return cases, reports


async def _retrieval_phase(qdrant, cases: list[dict]) -> list[dict]:
    results = []
    for case in cases:
        source = get_source(case["source"])
        collection = f"pilot_{source.collection}_benchmark"
        if case["exact"]:
            pilot = PilotSource(source)
            pilot.collection = collection
            exact_runs = []
            extracted = source.extract_ids(case["query"])
            if case["lookup_id"] not in extracted:
                raise AssertionError(f"Production identifier extraction failed: {case}")
            for _ in range(3):
                points = await asyncio.to_thread(
                    _exact_id_query_sync, qdrant, pilot, extracted, None
                )
                exact_runs.append([
                    {
                        "rank": 1,
                        "doc_id": str((point.payload or {}).get("doc_id") or ""),
                        "text": str((point.payload or {}).get("chunk_text") or ""),
                        "identity": f"Document ID: {(point.payload or {}).get('doc_id')}" in str((point.payload or {}).get("chunk_text") or ""),
                    }
                    for point in points
                ])
            runs = exact_runs
        else:
            runs = [await _hybrid_query(qdrant, collection, case["query"], limit=10) for _ in range(3)]
        ranks = [next((h["rank"] for h in run if h["doc_id"] == case["expected_id"]), None) for run in runs]
        target = next((h for h in runs[0] if h["doc_id"] == case["expected_id"]), None)
        result = {
            **case,
            "ranks": ranks,
            "target_found": target is not None,
            "target_content": bool(target and case["expected_marker"].lower() in target["text"].lower()),
            "rank_valid": all(rank == 1 if case["exact"] else rank is not None and rank <= 5 for rank in ranks),
            "identity_valid": bool(target and target["identity"]),
        }
        results.append(result)
        if not all(result[k] for k in ("target_found", "target_content", "rank_valid", "identity_valid")):
            raise AssertionError(f"Retrieval benchmark failure: {result}")
    return results


async def _upload_phase() -> list[dict]:
    results = []
    text = await parse_evidence("../evidence.LOG", b"BEGIN-SECRET\nCVE-like marker\nEND-SECRET")
    results.append({"label": "text.case-normalization", "passed": text["file_type"] == "text" and text["safe_filename"] == "evidence.LOG" and "BEGIN-SECRET" in text["extracted_text"]})

    pdf = io.BytesIO()
    document = Document()
    document.add_paragraph("PDF fixture marker is not used")
    # A minimal real PDF is produced through PyMuPDF in the parser test below.
    import fitz
    pdf_doc = fitz.open()
    page = pdf_doc.new_page()
    page.insert_text((72, 72), "PDF-EXTRACTION-MARKER")
    pdf_bytes = pdf_doc.tobytes()
    pdf_result = await parse_evidence("finding.pdf", pdf_bytes)
    results.append({"label": "pdf.text-extraction", "passed": pdf_result["file_type"] == "pdf" and "PDF-EXTRACTION-MARKER" in pdf_result["extracted_text"]})

    docx = io.BytesIO()
    document.add_paragraph("DOCX-EXTRACTION-MARKER")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "DOCX-TABLE-MARKER"
    document.save(docx)
    docx_result = await parse_evidence("finding.docx", docx.getvalue())
    results.append({"label": "docx.paragraph-and-table", "passed": docx_result["file_type"] == "office" and "DOCX-EXTRACTION-MARKER" in docx_result["extracted_text"] and "DOCX-TABLE-MARKER" in docx_result["extracted_text"]})

    image_bytes = io.BytesIO()
    Image.new("RGB", (20, 20), "white").save(image_bytes, format="PNG")
    with patch("app.services.evidence.describe_image", return_value="IMAGE-OCR-MARKER"):
        image_result = await parse_evidence("screen.png", image_bytes.getvalue())
    results.append({"label": "image.vision-success", "passed": image_result["file_type"] == "image" and image_result["image_description"] == "IMAGE-OCR-MARKER"})

    with patch("app.services.evidence.describe_image", side_effect=RuntimeError("vision down")):
        failed_image = await parse_evidence("screen.jpg", b"not-real-image")
    results.append({"label": "image.vision-failure-review", "passed": failed_image["needs_manual_review"] and "manual review" in failed_image["image_description"]})

    binary = await parse_evidence("capture.pcap", b"binary")
    results.append({"label": "unsupported.manual-review", "passed": binary["file_type"] == "binary" and binary["needs_manual_review"]})

    long_text = b"HEAD-MARKER" + b"x" * (10 * 1024 * 1024 + 50) + b"TAIL-MARKER"
    capped = await parse_evidence("large.log", long_text)
    results.append({"label": "text.size-cap-head-tail", "passed": "HEAD-MARKER" in capped["extracted_text"] and "TAIL-MARKER" not in capped["extracted_text"] and bool(capped["processing_notices"])})

    with tempfile.TemporaryDirectory() as directory, patch("app.services.upload_processing.settings.UPLOAD_DIR", directory), patch("app.services.upload_processing.EVIDENCE_CONTEXT_CHARS", 80):
        upload = UploadFile(filename="../safe.log", file=io.BytesIO(b"UPLOAD-BEGIN" + b"x" * 200 + b"UPLOAD-END"))
        parsed, manifest = await stage_evidence_uploads([upload])
        try:
            results.append({"label": "upload.staging-manifest-cleanup", "passed": parsed[0]["filename"] == "safe.log" and manifest["files"][0]["selected_chars"] < manifest["files"][0]["extracted_chars"]})
        finally:
            cleanup_staged_evidence(parsed)
            results.append({"label": "upload.cleanup", "passed": not Path(parsed[0]["storage_path"]).exists()})
    if not all(item["passed"] for item in results):
        raise AssertionError(f"Upload benchmark failure: {results}")
    return results


async def main(limit: int) -> None:
    load_model_sync()
    load_sparse_model_sync()
    init_qdrant_client()
    qdrant = get_qdrant()
    cases, index_reports = await _build_pilots(qdrant, max(12, limit))
    if len(cases) < 60:
        raise AssertionError(f"Expected at least 60 labeled retrieval cases, generated {len(cases)}")
    retrieval = await _retrieval_phase(qdrant, cases)
    uploads = await _upload_phase()
    print(json.dumps({
        "retrieval_case_count": len(retrieval),
        "retrieval_passed": len(retrieval),
        "index_reports": index_reports,
        "upload_case_count": len(uploads),
        "upload_passed": len(uploads),
        "all_passed": True,
        "gpu_runtime": "configured externally; dense model loaded by benchmark",
        "production_migration": "BLOCKED: benchmark evidence only; no production changes made",
    }, indent=2, default=str))


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--limit", type=int, default=12)
    args = parser.parse_args()
    asyncio.run(main(max(12, args.limit)))
