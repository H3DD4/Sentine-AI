"""
Report generation service.
Generates branded DOCX reports from analyst-approved finding data.
"""

import io
from docx import Document
from datetime import datetime, timezone
from app.models import Finding

DEFAULT_SECTIONS = {
    "executive_summary", "scope_methodology", "findings_overview",
    "detailed_findings", "attack_mapping", "evidence_gaps", "disclaimer",
}

def _verdict_value(finding: Finding) -> str:
    verdict = finding.verdict
    return verdict.value if hasattr(verdict, "value") else str(verdict or "")


def _add_text_blocks(doc: Document, text: str) -> None:
    for block in (part.strip() for part in text.split("\n\n")):
        if block:
            doc.add_paragraph(block)


async def generate_report_docx(
    findings: list[Finding],
    engagement_title: str,
    client_name: str,
    template_path: str | None = None,
    sections: set[str] | None = None,
) -> bytes:
    """
    Generates a branded DOCX report.
    Every section is assembled from analyst-approved fields. Export does not
    call an LLM, so model instructions or invented prose cannot leak into the
    client deliverable.
    Returns raw bytes of the .docx file.
    """
    sections = sections or DEFAULT_SECTIONS
    doc = Document(template_path) if template_path else Document()

    if template_path:
        _replace_placeholders(doc, {
            "{{CLIENT_NAME}}": client_name,
            "{{ENGAGEMENT_TITLE}}": engagement_title,
            "{{GENERATED_AT}}": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
            "{{TOTAL_FINDINGS}}": str(len(findings)),
        })
        doc.add_page_break()

    if not template_path:
        doc.add_heading(f"{client_name} - Red Team Findings Report", 0)
        doc.add_paragraph(f"Engagement: {engagement_title}")
        doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph(f"Total findings: {len(findings)}")
        doc.add_paragraph("CONFIDENTIAL - TLP:RED - For authorised recipients only.")
        doc.add_page_break()

    confirmed = sum(1 for finding in findings if _verdict_value(finding) == "confirmed")
    likely = sum(1 for finding in findings if _verdict_value(finding) == "likely")
    pending = sum(1 for finding in findings if not _verdict_value(finding))
    insufficient = sum(1 for finding in findings if _verdict_value(finding) == "insufficient")
    false_positives = sum(1 for finding in findings if _verdict_value(finding) == "false_positive")

    # Assemble document
    if "executive_summary" in sections:
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph(
            f"This report presents the results of the {engagement_title} conducted for {client_name}. "
            f"It documents {len(findings)} finding{'s' if len(findings) != 1 else ''} selected by the "
            "analyst for inclusion in this deliverable."
        )
        doc.add_paragraph(
            f"The selected findings comprise {confirmed} confirmed, {likely} likely, "
            f"{insufficient} insufficient, and {false_positives} false positive finding(s). "
            "Each detailed finding preserves the analyst-approved technical evidence, observed "
            "impact, validation rationale, and references available at the time of export."
        )
        doc.add_page_break()

    if "scope_methodology" in sections:
        doc.add_heading("Scope and Methodology", 1)
        doc.add_paragraph(
            "This deliverable is limited to the targets, components, endpoints, and evidence explicitly "
            "recorded in the selected findings. Conclusions are based on observed behaviour and supplied "
            "technical evidence; unrecorded assumptions are not presented as established facts."
        )

    if "findings_overview" in sections:
        doc.add_heading("Findings Overview", 1)
        overview = doc.add_table(rows=1, cols=4)
        overview.style = "Table Grid"
        for cell, label in zip(overview.rows[0].cells, ("Finding", "Verdict", "Severity", "CVSS")):
            cell.text = label
        for finding in findings:
            row = overview.add_row().cells
            verdict = _verdict_value(finding)
            row[0].text = finding.title
            row[1].text = verdict.replace("_", " ").title()
            row[2].text = finding.severity
            row[3].text = f"{finding.cvss_score:.1f}" if finding.cvss_score is not None else "N/A"
        doc.add_page_break()

    if "detailed_findings" in sections:
        doc.add_heading("Detailed Findings", 1)
    for i, finding in enumerate(findings, 1) if "detailed_findings" in sections else []:
        doc.add_heading(f"{i}. {finding.title}", 2)

        # Data table — programmatically inserted, NOT LLM generated
        table = doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"
        rows = table.rows
        rows[0].cells[0].text = "Verdict"
        verdict = _verdict_value(finding)
        rows[0].cells[1].text = verdict.replace("_", " ").title() if verdict else "Pending"
        rows[1].cells[0].text = "Severity"
        rows[1].cells[1].text = finding.severity
        rows[2].cells[0].text = "Confidence"
        rows[2].cells[1].text = f"{int((finding.confidence or 0) * 100)}%"
        rows[3].cells[0].text = "Matched CVEs"
        rows[3].cells[1].text = ", ".join(finding.matched_cves) if finding.matched_cves else "None"
        rows[4].cells[0].text = "CVSS Base Score"
        cvss_score = getattr(finding, "cvss_score", None)
        rows[4].cells[1].text = f"{cvss_score:.1f}" if cvss_score is not None else "N/A"
        rows[5].cells[0].text = "CVSS Vector"
        rows[5].cells[1].text = finding.cvss_vector or "N/A"

        doc.add_paragraph()
        doc.add_heading("Description", 3)
        _add_text_blocks(doc, finding.description)

        doc.add_heading("Affected Scope", 3)
        _add_text_blocks(doc, finding.affected_scope)

        doc.add_heading("Technical evidence", 3)
        _add_text_blocks(doc, finding.technical_evidence)

        if finding.reproduction_steps:
            doc.add_heading("Reproduction Steps", 3)
            for step in finding.reproduction_steps:
                doc.add_paragraph(step, style="List Number")

        doc.add_heading("Impact", 3)
        _add_text_blocks(doc, finding.impact)

        doc.add_heading("Analysis", 3)
        doc.add_paragraph(finding.reasoning or finding.impact)

        if finding.matched_techniques and "attack_mapping" in sections:
            doc.add_heading("MITRE ATT&CK Techniques", 3)
            doc.add_paragraph(", ".join(finding.matched_techniques))

        if finding.missing_evidence and "evidence_gaps" in sections:
            doc.add_heading("Evidence Gaps", 3)
            for gap in finding.missing_evidence:
                doc.add_paragraph(gap, style="List Bullet")

        doc.add_page_break()

    # Disclaimer
    if "disclaimer" in sections:
        doc.add_heading("Disclaimer", 1)
        doc.add_paragraph(
            "This report was generated for authorized penetration testing purposes only. "
            "All findings were identified during a sanctioned engagement. "
            "Reproduction or distribution outside the authorized scope is prohibited."
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _replace_placeholders(doc: Document, replacements: dict[str, str]) -> None:
    containers = [doc.paragraphs]
    containers.extend(
        cell.paragraphs for table in doc.tables for row in table.rows for cell in row.cells
    )
    for section in doc.sections:
        containers.extend((section.header.paragraphs, section.footer.paragraphs))
    for paragraphs in containers:
        for paragraph in paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    replaced = paragraph.text.replace(placeholder, value)
                    if paragraph.runs:
                        paragraph.runs[0].text = replaced
                        for run in paragraph.runs[1:]:
                            run.text = ""
                    else:
                        paragraph.add_run(replaced)
