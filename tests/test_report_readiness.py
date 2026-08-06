import asyncio
import io
import unittest
from unittest.mock import AsyncMock, patch

from docx import Document

from app.schemas import ChatMessage, ReportDraft
from app.services.report import generate_report_docx
from app.routers.report import _missing_report_fields
from app.kb.base import SearchOutcome
from app.routers.chat import (
    _bounded_provider_messages,
    _conversation_outcome,
    _needs_retrieval,
    _social_reply,
)
from app.services.report_readiness import (
    READINESS_THRESHOLD,
    _normalize_extraction,
    assess_conversation,
    draft_to_finding,
    score_report_draft,
)


class ReportReadinessTests(unittest.TestCase):
    def test_long_assessment_is_preserved_in_provider_context(self):
        assessment = "Observed command execution evidence. " * 80
        messages = [ChatMessage(role="user", content=assessment)]

        bounded = _bounded_provider_messages(messages)

        self.assertGreater(len(assessment), 1500)
        self.assertEqual(bounded, [{"role": "user", "content": assessment}])

    def test_provider_context_is_bounded_and_keeps_latest_user_turn(self):
        latest = "Latest evidence must remain intact."
        messages = [
            ChatMessage(role="user", content=f"old user {index} " + "x" * 5000)
            if index % 2 == 0
            else ChatMessage(role="assistant", content=f"old assistant {index} " + "y" * 5000)
            for index in range(20)
        ]
        messages.append(ChatMessage(role="user", content=latest))

        bounded = _bounded_provider_messages(messages, max_messages=8, max_chars=12_000)

        self.assertEqual(bounded[-1], {"role": "user", "content": latest})
        self.assertLessEqual(len(bounded), 8)
        self.assertLessEqual(sum(len(message["content"]) for message in bounded), 12_000)
        self.assertEqual(bounded[0]["role"], "user")

    def test_retrieval_failure_degrades_instead_of_raising(self):
        with patch("app.routers.chat.federated_search", new=AsyncMock(side_effect=RuntimeError("model unavailable"))):
            outcome = asyncio.run(_conversation_outcome("Analyze this finding", [], AsyncMock()))

        self.assertIsInstance(outcome, SearchOutcome)
        self.assertTrue(outcome.degraded)
        self.assertIn("not grounded", outcome.notes[0])

    def test_greeting_does_not_trigger_knowledge_retrieval(self):
        self.assertFalse(_needs_retrieval("hello"))
        self.assertFalse(_needs_retrieval("Good morning!"))
        self.assertTrue(_needs_retrieval("Is Struts 2.3.5 affected by CVE-2017-5638?"))
        self.assertIsNotNone(_social_reply("hello"))
        self.assertIsNone(_social_reply("Analyze this Struts finding"))

    def test_greeting_does_not_call_llm_for_readiness(self):
        with patch("app.services.report_readiness.AsyncLLMClient") as client:
            result = asyncio.run(
                assess_conversation([ChatMessage(role="user", content="Hello!")])
            )

        self.assertEqual(result.score, 0)
        client.assert_not_called()

    def test_empty_draft_is_not_eligible(self):
        result = score_report_draft(ReportDraft())

        self.assertEqual(result.score, 0)
        self.assertFalse(result.eligible)
        self.assertEqual(result.status, "not_ready")

    def test_minimum_technical_draft_is_eligible_below_full_score(self):
        result = score_report_draft(
            ReportDraft(
                title="SQL injection",
                description="The id parameter appears injectable.",
                technical_evidence="A quote causes a database error and sqlmap confirms injection.",
            )
        )

        self.assertGreaterEqual(result.score, READINESS_THRESHOLD)
        self.assertLess(result.score, 10)
        self.assertTrue(result.eligible)
        self.assertEqual(result.status, "reportable")

    def test_score_alone_cannot_replace_required_technical_substance(self):
        result = score_report_draft(
            ReportDraft(
                title="Potential access control issue",
                description="A possible authorization weakness was discussed.",
                affected_scope="Customer portal",
                impact="Another user's records may be exposed.",
                remediation=["Enforce object authorization."],
                matched_techniques=["T1190"],
                verdict="likely",
            )
        )

        self.assertGreaterEqual(result.score, READINESS_THRESHOLD)
        self.assertFalse(result.eligible)
        self.assertIn("technical evidence", result.summary.lower())

    def test_complete_draft_reaches_ready_state(self):
        result = score_report_draft(
            ReportDraft(
                title="SQL injection",
                affected_scope="GET /users?id= on app.example",
                description="The endpoint is vulnerable to unauthenticated SQL injection.",
                technical_evidence="sqlmap confirmed injection and returned the database banner.",
                reproduction_steps=["Send a quote in id.", "Confirm with sqlmap."],
                impact="An attacker can read and modify application data.",
                severity="Critical",
                cvss_score=9.8,
                cvss_vector="CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H",
                remediation=["Use parameterized queries."],
                matched_cves=["CVE-2024-0001"],
                matched_techniques=["T1190"],
                verdict="confirmed",
                confidence=0.94,
            )
        )

        self.assertEqual(result.score, 10)
        self.assertTrue(result.eligible)
        self.assertEqual(result.status, "ready")
        self.assertEqual(result.missing, [])

    def test_confirmed_struts_rce_evidence_is_reportable(self):
        result = score_report_draft(
            ReportDraft(
                title="Remote code execution via Struts OGNL injection",
                affected_scope="POST /struts2-showcase/fileupload/doUpload.action",
                description="A crafted multipart Content-Type header executes operating-system commands.",
                technical_evidence="HTTP 200 returned uid=0(root) gid=0(root) after the id command.",
                reproduction_steps=[
                    "Send the crafted multipart request to doUpload.action.",
                    "Observe the id command output in the response.",
                ],
                impact="An unauthenticated attacker can execute commands as root.",
            )
        )

        self.assertEqual(result.score, 6.5)
        self.assertTrue(result.eligible)
        self.assertIn("Impact and risk rating", result.missing)

    def test_malformed_extracted_verdict_is_ignored(self):
        data = _normalize_extraction(
            {
                "title": "Struts remote code execution",
                "verdict": "Finding confirmed: Remote Code Execution in multipart parsing.",
                "confidence": "95%",
            }
        )

        self.assertIsNone(data["verdict"])
        self.assertEqual(data["confidence"], 0.95)
        self.assertEqual(data["title"], "Struts remote code execution")

    def test_assessment_keeps_draft_when_llm_returns_malformed_verdict(self):
        extracted = {
            "title": "Remote code execution via Struts OGNL injection",
            "affected_scope": "POST /struts2-showcase/fileupload/doUpload.action",
            "description": "A crafted multipart Content-Type header executes commands.",
            "technical_evidence": "HTTP 200 returned uid=0(root).",
            "reproduction_steps": ["Send the crafted multipart request."],
            "impact": "Remote command execution as root.",
            "remediation": [],
            "matched_cves": [],
            "matched_techniques": [],
            "verdict": "Finding confirmed: Remote Code Execution in multipart parsing.",
            "confidence": "95%",
        }
        client = AsyncMock()
        client.generate_validation.return_value = extracted

        with patch("app.services.report_readiness.AsyncLLMClient", return_value=client):
            result = asyncio.run(
                assess_conversation(
                    [ChatMessage(role="user", content="Struts multipart RCE returned uid=0(root).")]
                )
            )

        self.assertTrue(result.eligible)
        self.assertIsNone(result.draft.verdict)
        self.assertEqual(result.draft.confidence, 0.95)

    def test_assessment_falls_back_when_llm_is_unavailable(self):
        client = AsyncMock()
        client.generate_validation.side_effect = RuntimeError("provider quota exceeded")

        with patch("app.services.report_readiness.AsyncLLMClient", return_value=client):
            result = asyncio.run(
                assess_conversation(
                    [
                        ChatMessage(
                            role="user",
                            content="SQL injection in POST /login. A quote returns HTTP 500.",
                        )
                    ]
                )
            )

        self.assertEqual(result.draft.title, "SQL injection in POST /login")
        self.assertIn("HTTP 500", result.draft.description)
        self.assertEqual(result.draft.technical_evidence, "")
        self.assertFalse(result.eligible)
        self.assertIsNotNone(result.assessment_notice)
        self.assertIn("conservative", result.summary.lower())

    def test_assessment_falls_back_when_extracted_shape_is_invalid(self):
        client = AsyncMock()
        client.generate_validation.return_value = {
            "title": ["invalid", "title"],
            "reproduction_steps": "not a list",
        }

        with patch("app.services.report_readiness.AsyncLLMClient", return_value=client):
            result = asyncio.run(
                assess_conversation(
                    [ChatMessage(role="user", content="Potential SQL injection in POST /login.")]
                )
            )

        self.assertEqual(result.draft.title, "Potential SQL injection in POST /login")
        self.assertIsNotNone(result.assessment_notice)

    def test_conversation_draft_generates_a_docx(self):
        draft = ReportDraft(
            title="SQL injection",
            description="The id parameter is injectable.",
            technical_evidence="sqlmap confirmed boolean-based injection.",
            reproduction_steps=["Send a quote in id.", "Confirm with sqlmap."],
            impact="An attacker can disclose database records.",
            remediation=["Use parameterized queries."],
        )

        content = asyncio.run(
            generate_report_docx(
                [draft_to_finding(draft)],
                "External penetration test",
                "Northwind",
            )
        )

        self.assertTrue(content.startswith(b"PK"))
        self.assertGreater(len(content), 1_000)
        document = Document(io.BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        report_text = f"{text}\n{table_text}"
        self.assertIn("Northwind", report_text)
        self.assertIn("Technical evidence", report_text)
        self.assertNotIn("Write a 3-paragraph", report_text)
        self.assertNotIn("VerdictEnum.", report_text)
        self.assertNotIn("Recommendations", report_text)

    def test_incomplete_finding_is_blocked_from_client_export(self):
        finding = draft_to_finding(
            ReportDraft(title="Command injection", description="Commands execute through host.")
        )

        gaps = _missing_report_fields(finding)

        self.assertIn("technical evidence", gaps)
        self.assertIn("impact", gaps)
        self.assertIn("severity", gaps)
        self.assertIn("analyst verdict", gaps)

    def test_uploaded_template_is_used_and_placeholders_are_replaced(self):
        template = Document()
        paragraph = template.add_paragraph()
        paragraph.add_run("{{CLIENT_")
        paragraph.add_run("NAME}}")
        paragraph.add_run(" / {{ENGAGEMENT_TITLE}}")
        with __import__("tempfile").TemporaryDirectory() as directory:
            path = __import__("pathlib").Path(directory) / "template.docx"
            template.save(path)
            draft = ReportDraft(
                title="Command injection",
                affected_scope="POST /ping",
                description="The host parameter executes operating-system commands.",
                technical_evidence="The response returned uid=33(www-data).",
                impact="An attacker can execute commands on the host.",
                severity="High",
                verdict="confirmed",
            )

            content = asyncio.run(
                generate_report_docx(
                    [draft_to_finding(draft)],
                    "External test",
                    "Northwind",
                    template_path=str(path),
                    sections={"detailed_findings"},
                )
            )

        document = Document(io.BytesIO(content))
        report_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("Northwind / External test", report_text)
        self.assertNotIn("{{CLIENT_NAME}}", report_text)
        self.assertNotIn("Executive Summary", report_text)


if __name__ == "__main__":
    unittest.main()
